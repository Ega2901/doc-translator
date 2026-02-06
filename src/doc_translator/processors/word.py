import copy
import re
import warnings
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from doc_translator.models.chunk import Chunk, ElementMetadata, ElementType, TranslatedChunk
from doc_translator.processors.base import DocumentProcessor


class WordDocumentProcessor(DocumentProcessor):
    def get_supported_extensions(self) -> list[str]:
        return [".docx"]

    def load(self, path: str | Path) -> DocumentType:
        return Document(str(path))

    def save(self, document: DocumentType, path: str | Path) -> None:
        document.save(str(path))

    def _extract_paragraph_metadata(self, paragraph: Paragraph) -> ElementMetadata:
        style_name = paragraph.style.name if paragraph.style else None
        font_name = None
        font_size = None
        bold = False
        italic = False
        underline = False

        if paragraph.runs:
            run = paragraph.runs[0]
            font = run.font
            font_name = font.name
            font_size = font.size.pt if font.size else None
            bold = font.bold or False
            italic = font.italic or False
            underline = font.underline or False

        alignment = None
        if paragraph.alignment is not None:
            alignment = str(paragraph.alignment)

        return ElementMetadata(
            element_type=ElementType.PARAGRAPH,
            style_name=style_name,
            font_name=font_name,
            font_size=font_size,
            bold=bold,
            italic=italic,
            underline=underline,
            alignment=alignment,
        )

    def _extract_table_metadata(self, table: Table) -> ElementMetadata:
        return ElementMetadata(
            element_type=ElementType.TABLE,
            rows=len(table.rows),
            cols=len(table.columns),
        )

    def _get_table_text(self, table: Table) -> str:
        rows_text = []
        for row in table.rows:
            cells_text = []
            for cell in row.cells:
                cells_text.append(cell.text.strip())
            rows_text.append(" | ".join(cells_text))
        return "\n".join(rows_text)

    def _iter_block_items(self, document: DocumentType):
        body = document.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, document)
            elif child.tag == qn("w:tbl"):
                yield Table(child, document)

    def chunk(self, path: str | Path) -> list[Chunk]:
        document = self.load(path)
        chunks: list[Chunk] = []

        current_texts: list[str] = []
        current_metadata: list[ElementMetadata] = []
        current_elements: list[Any] = []
        current_char_count = 0
        chunk_index = 0

        def flush_chunk():
            nonlocal chunk_index, current_texts, current_metadata, current_elements, current_char_count

            if current_texts:
                text = "\n\n".join(current_texts)
                chunks.append(
                    Chunk(
                        index=chunk_index,
                        text=text,
                        char_count=len(text),
                        metadata=current_metadata.copy(),
                        original_elements=current_elements.copy(),
                        source_file=str(path),
                    )
                )
                chunk_index += 1

            current_texts = []
            current_metadata = []
            current_elements = []
            current_char_count = 0

        for element in self._iter_block_items(document):
            if isinstance(element, Paragraph):
                text = element.text.strip()
                if not text:
                    continue

                text_len = len(text)
                separator_len = 2 if current_texts else 0

                if current_char_count + separator_len + text_len > self.max_chars and current_texts:
                    flush_chunk()
                    separator_len = 0

                current_texts.append(text)
                current_metadata.append(self._extract_paragraph_metadata(element))
                current_elements.append(("paragraph", element._element))
                current_char_count += separator_len + text_len

            elif isinstance(element, Table):
                table_text = self._get_table_text(element)
                full_table_text = f"[ТАБЛИЦА]\n{table_text}\n[/ТАБЛИЦА]"
                table_len = len(full_table_text)
                separator_len = 2 if current_texts else 0

                if table_len > self.max_chars:
                    warnings.warn(
                        f"Таблица ({table_len} символов) превышает max_chars ({self.max_chars}). "
                        "Таблица будет добавлена отдельным чанком."
                    )

                if current_char_count + separator_len + table_len > self.max_chars and current_texts:
                    flush_chunk()
                    separator_len = 0

                current_texts.append(full_table_text)
                current_metadata.append(self._extract_table_metadata(element))
                current_elements.append(("table", element._tbl))
                current_char_count += separator_len + table_len

                if table_len > self.max_chars:
                    flush_chunk()

        flush_chunk()

        return chunks

    def concatenate(
        self,
        translated_chunks: list[TranslatedChunk],
        output_path: str | Path,
    ) -> None:
        if not translated_chunks:
            raise ValueError("Список переведённых чанков пуст")

        sorted_chunks = sorted(translated_chunks, key=lambda c: c.index)

        source_file = sorted_chunks[0].original.source_file
        if source_file:
            template_doc = self.load(source_file)
        else:
            template_doc = Document()

        new_doc = Document()
        self._copy_styles(template_doc, new_doc)

        for chunk in sorted_chunks:
            translated_lines = chunk.translated_text.split("\n\n")
            line_idx = 0

            for i, (elem_type, _) in enumerate(chunk.original_elements):
                if line_idx >= len(translated_lines):
                    break

                metadata = chunk.metadata[i] if i < len(chunk.metadata) else None

                if elem_type == "paragraph":
                    text = translated_lines[line_idx]
                    line_idx += 1

                    para = new_doc.add_paragraph()

                    if metadata and metadata.style_name:
                        try:
                            para.style = metadata.style_name
                        except KeyError:
                            pass

                    run = para.add_run(text)
                    if metadata:
                        if metadata.bold:
                            run.bold = True
                        if metadata.italic:
                            run.italic = True
                        if metadata.underline:
                            run.underline = True
                        if metadata.font_name:
                            run.font.name = metadata.font_name
                        if metadata.font_size:
                            run.font.size = int(metadata.font_size * 12700)

                elif elem_type == "table":
                    # Собираем все блоки, относящиеся к этой таблице (модель могла вставить \n\n внутри)
                    table_blocks: list[str] = []
                    expected_rows = metadata.rows if metadata else None
                    first_block = None

                    while line_idx < len(translated_lines):
                        block = translated_lines[line_idx]
                        if first_block is None:
                            first_block = block
                        table_blocks.append(block)
                        line_idx += 1
                        combined = "\n".join(table_blocks)
                        if "[/ТАБЛИЦА]" in combined:
                            break
                        if expected_rows is not None:
                            raw = combined.replace("[ТАБЛИЦА]", "").replace("[/ТАБЛИЦА]", "").strip()
                            lines = [ln for ln in raw.split("\n") if ln.strip()]
                            if len(lines) >= expected_rows:
                                break
                        # Таблица без маркеров: один блок = одна таблица, не тянем следующие
                        if "[ТАБЛИЦА]" not in (first_block or "") and "[ТАБЛИЦА]" not in combined:
                            break
                    table_text = "\n".join(table_blocks)

                    table_text = table_text.replace("[ТАБЛИЦА]", "").replace("[/ТАБЛИЦА]", "").strip()
                    # Разбор строк таблицы: допускаем " | ", "|", " |" и т.д.; пустые строки пропускаем
                    rows_data = [
                        [s.strip() for s in re.split(r"\s*\|\s*", line)]
                        for line in table_text.split("\n")
                        if line.strip()
                    ]
                    flat_cells = [c for row in rows_data for c in row]

                    tbl_element = chunk.original_elements[i][1]
                    cloned_tbl = copy.deepcopy(tbl_element)
                    new_doc.element.body.append(cloned_tbl)
                    table = Table(cloned_tbl, new_doc)

                    cell_list = []
                    for row in table.rows:
                        for cell in row.cells:
                            cell_list.append(cell)
                    for idx, cell in enumerate(cell_list):
                        if idx < len(flat_cells):
                            cell.text = flat_cells[idx]
                        else:
                            cell.text = ""

        self.save(new_doc, output_path)

    def _copy_styles(self, source: DocumentType, target: DocumentType) -> None:
        pass
